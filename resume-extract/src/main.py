import asyncio
import json
import logging
import os
import re
import shutil
import socket
import time
from collections import Counter
from contextlib import asynccontextmanager, suppress
from datetime import datetime, timedelta, timezone
from hmac import compare_digest
from io import BytesIO
from typing import Any
from urllib.parse import quote, urlencode, urlparse, urlunparse
from uuid import UUID

import httpx
import sentry_sdk
import websockets
from dotenv import load_dotenv
from fastapi import FastAPI, Header
from fastapi.responses import JSONResponse
from pypdf import PdfReader
from pydantic import BaseModel
from docx import Document
from sentry_sdk.integrations.fastapi import FastApiIntegration

load_dotenv()

def _env_flag(name: str, default: bool = False) -> bool:
    raw = os.getenv(name)
    if raw is None:
        return default
    return raw.strip().lower() in {"1", "true", "yes", "on"}


SENTRY_DSN = os.getenv("SENTRY_DSN", "").strip()
if SENTRY_DSN:
    sentry_sdk.init(
        dsn=SENTRY_DSN,
        environment=os.getenv("SENTRY_ENVIRONMENT"),
        release=os.getenv("SENTRY_RELEASE"),
        traces_sample_rate=float(os.getenv("SENTRY_TRACES_SAMPLE_RATE", "0")),
        profiles_sample_rate=float(os.getenv("SENTRY_PROFILES_SAMPLE_RATE", "0")),
        send_default_pii=_env_flag("SENTRY_SEND_DEFAULT_PII", default=False),
        integrations=[FastApiIntegration()],
    )

SUPABASE_URL = os.getenv("SUPABASE_URL")
SUPABASE_SERVICE_ROLE_KEY = os.getenv("SUPABASE_SERVICE_ROLE_KEY")
SUPABASE_DB_SCHEMA = os.getenv("SUPABASE_DB_SCHEMA", "public")
SUPABASE_STORAGE_BUCKET = os.getenv("SUPABASE_STORAGE_BUCKET")
EXTRACT_AUTH_TOKEN = os.getenv("EXTRACT_AUTH_TOKEN")
EXTRACT_DOWNLOAD_TIMEOUT_SECONDS = float(os.getenv("EXTRACT_DOWNLOAD_TIMEOUT_SECONDS", "60"))
MAX_RESUME_FILE_SIZE_MB = int(os.getenv("MAX_RESUME_FILE_SIZE_MB", "15"))
WORKER_POLL_INTERVAL_SECONDS = float(os.getenv("WORKER_POLL_INTERVAL_SECONDS", "2"))
WORKER_REALTIME_ENABLED = _env_flag("WORKER_REALTIME_ENABLED", default=True)
WORKER_REALTIME_HEARTBEAT_SECONDS = float(os.getenv("WORKER_REALTIME_HEARTBEAT_SECONDS", "25"))
WORKER_REALTIME_RECONNECT_SECONDS = float(os.getenv("WORKER_REALTIME_RECONNECT_SECONDS", "5"))
WORKER_ATTACH_TO_API = _env_flag("WORKER_ATTACH_TO_API", default=False)
WORKER_CLAIM_OWNER = os.getenv("WORKER_CLAIM_OWNER", "").strip() or f"worker:{socket.gethostname()}:{os.getpid()}"
API_CLAIM_OWNER = os.getenv("API_CLAIM_OWNER", "").strip() or f"api:{socket.gethostname()}:{os.getpid()}"
WORKER_LEASE_SECONDS = int(os.getenv("WORKER_LEASE_SECONDS", "120"))
RUN_HEARTBEAT_INTERVAL_SECONDS = float(os.getenv("RUN_HEARTBEAT_INTERVAL_SECONDS", "15"))
WORKER_MAX_ATTEMPTS = int(os.getenv("WORKER_MAX_ATTEMPTS", "3"))
WORKER_RETRY_DELAY_SECONDS = float(os.getenv("WORKER_RETRY_DELAY_SECONDS", "15"))
WORKER_QUEUE_SNAPSHOT_INTERVAL_SECONDS = float(os.getenv("WORKER_QUEUE_SNAPSHOT_INTERVAL_SECONDS", "60"))
WORKER_HEARTBEAT_INTERVAL_SECONDS = float(os.getenv("WORKER_HEARTBEAT_INTERVAL_SECONDS", "15"))
WORKER_ACTIVE_WINDOW_SECONDS = int(os.getenv("WORKER_ACTIVE_WINDOW_SECONDS", "45"))
WORKER_STALE_RESET_LIMIT = int(os.getenv("WORKER_STALE_RESET_LIMIT", "100"))
WORKER_STALE_RESET_INTERVAL_SECONDS = float(os.getenv("WORKER_STALE_RESET_INTERVAL_SECONDS", "60"))
GENERATION_WORKER_ENABLED = _env_flag("GENERATION_WORKER_ENABLED", default=True)
GENERATION_WORKER_CLAIM_OWNER = os.getenv("GENERATION_WORKER_CLAIM_OWNER", "").strip() or f"generate:{socket.gethostname()}:{os.getpid()}"
GENERATION_WORKER_LEASE_SECONDS = int(os.getenv("GENERATION_WORKER_LEASE_SECONDS", "300"))
GENERATION_WORKER_MAX_ATTEMPTS = int(os.getenv("GENERATION_WORKER_MAX_ATTEMPTS", "3"))
GENERATION_RETRY_DELAY_SECONDS = float(os.getenv("GENERATION_RETRY_DELAY_SECONDS", "30"))
GENERATION_RUN_HEARTBEAT_INTERVAL_SECONDS = float(os.getenv("GENERATION_RUN_HEARTBEAT_INTERVAL_SECONDS", "20"))
PDF_WORKER_ENABLED = _env_flag("PDF_WORKER_ENABLED", default=True)
PDF_WORKER_CLAIM_OWNER = os.getenv("PDF_WORKER_CLAIM_OWNER", "").strip() or f"pdf:{socket.gethostname()}:{os.getpid()}"
PDF_WORKER_LEASE_SECONDS = int(os.getenv("PDF_WORKER_LEASE_SECONDS", "300"))
PDF_WORKER_MAX_ATTEMPTS = int(os.getenv("PDF_WORKER_MAX_ATTEMPTS", "3"))
PDF_RETRY_DELAY_SECONDS = float(os.getenv("PDF_RETRY_DELAY_SECONDS", "30"))
PDF_RUN_HEARTBEAT_INTERVAL_SECONDS = float(os.getenv("PDF_RUN_HEARTBEAT_INTERVAL_SECONDS", "20"))

MAX_RESUME_FILE_SIZE_BYTES = MAX_RESUME_FILE_SIZE_MB * 1024 * 1024
ALLOWED_BINARY_CONTENT_TYPES = {"application/octet-stream", "binary/octet-stream"}
ALLOWED_PDF_CONTENT_TYPES = {
    "application/pdf",
    "application/x-pdf",
    *ALLOWED_BINARY_CONTENT_TYPES,
}
ALLOWED_DOCX_CONTENT_TYPES = {
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "application/zip",
    *ALLOWED_BINARY_CONTENT_TYPES,
}

@asynccontextmanager
async def _app_lifespan(_: FastAPI):
    worker_task: asyncio.Task[None] | None = None
    if WORKER_ATTACH_TO_API:
        _log_event(phase="api_start_worker_attached")
        worker_task = asyncio.create_task(run_worker_forever())
    try:
        yield
    finally:
        if worker_task:
            worker_task.cancel()
            with suppress(asyncio.CancelledError):
                await worker_task


app = FastAPI(lifespan=_app_lifespan)
logger = logging.getLogger("resume_extract_service")
if not logger.handlers:
    logging.basicConfig(level=logging.INFO)

# Canonical error codes for extraction pipeline.
ERR_MISSING_CONFIG = "missing_config"
ERR_UNAUTHORIZED = "unauthorized"
ERR_DB_ERROR = "db_error"
ERR_RUN_NOT_QUEUED = "run_not_queued"
ERR_RUN_NOT_EXTRACTING = "run_not_extracting"
ERR_INVALID_RUN_ROW = "invalid_run_row"
ERR_DOWNLOAD_FAILED = "download_failed"
ERR_FILE_TOO_LARGE = "file_too_large"
ERR_UNSUPPORTED_CONTENT_TYPE = "unsupported_content_type"
ERR_EMPTY_RESUME_FILE = "empty_resume_file"
ERR_UNSUPPORTED_FILE_TYPE = "unsupported_file_type"
ERR_EMPTY_EXTRACTED_TEXT = "empty_extracted_text"
ERR_PARSE_ERROR = "parse_error"
ERR_OCR_UNAVAILABLE = "ocr_unavailable"
ERR_OCR_DEPENDENCY_MISSING = "ocr_dependency_missing"
ERR_GENERATE_NOT_QUEUED = "generate_not_queued"
ERR_GENERATE_NOT_CLAIMED = "generate_not_claimed"
ERR_FUNCTION_CALL_FAILED = "function_call_failed"
ERR_PDF_NOT_CLAIMED = "pdf_not_claimed"

BULLET_PREFIX_RE = re.compile(r"^[\-\*\u2022\u2023\u25E6\u2043\u2219\u00B7]+[\s\t]*")
PAGE_MARKER_RE = re.compile(
    r"^(page\s+\d+(\s+of\s+\d+)?|\d+\s*/\s*\d+|\d+\s+of\s+\d+)$",
    flags=re.IGNORECASE,
)
ZERO_WIDTH_CHARS_RE = re.compile(r"[\u200b\u200c\u200d\ufeff]")


class HttpError(Exception):
    def __init__(self, status: int, code: str, message: str) -> None:
        super().__init__(message)
        self.status = status
        self.code = code
        self.message = message


class ExtractRequest(BaseModel):
    run_id: UUID


def _is_supabase_opaque_key(value: str) -> bool:
    return value.startswith("sb_secret_") or value.startswith("sb_publishable_")


def _capture_exception(
    error: Exception,
    *,
    phase: str,
    run_id: UUID | None = None,
    parser: str | None = None,
) -> None:
    # Keep Sentry signal clean: report only unexpected errors and server-side HttpErrors.
    if isinstance(error, HttpError) and error.status < 500:
        return
    with sentry_sdk.push_scope() as scope:
        scope.set_tag("phase", phase)
        if isinstance(error, HttpError):
            scope.set_tag("error_code", error.code)
            scope.set_extra("http_status", error.status)
        if run_id is not None:
            scope.set_tag("run_id", str(run_id))
        if parser:
            scope.set_tag("parser", parser)
        sentry_sdk.capture_exception(error)


def _log_event(
    phase: str,
    run_id: UUID | None = None,
    parser: str | None = None,
    duration_ms: int | None = None,
    error_code: str | None = None,
    **extra: Any,
) -> None:
    payload: dict[str, Any] = {"event": "resume_extract", "phase": phase}
    if run_id is not None:
        payload["run_id"] = str(run_id)
    if parser:
        payload["parser"] = parser
    if duration_ms is not None:
        payload["duration_ms"] = duration_ms
    if error_code:
        payload["error_code"] = error_code
    payload.update(extra)
    logger.info(json.dumps(payload, separators=(",", ":"), default=str))


def _assert_extract_auth(token: str | None) -> None:
    if not EXTRACT_AUTH_TOKEN:
        raise HttpError(
            status=500,
            code=ERR_MISSING_CONFIG,
            message="EXTRACT_AUTH_TOKEN must be configured",
        )
    if not token or not compare_digest(token, EXTRACT_AUTH_TOKEN):
        raise HttpError(
            status=401,
            code=ERR_UNAUTHORIZED,
            message="Invalid extract auth token",
        )


def _assert_config() -> None:
    if not SUPABASE_URL or not SUPABASE_SERVICE_ROLE_KEY or not SUPABASE_STORAGE_BUCKET:
        raise HttpError(
            status=500,
            code=ERR_MISSING_CONFIG,
            message="SUPABASE_URL, SUPABASE_SERVICE_ROLE_KEY, and SUPABASE_STORAGE_BUCKET must be configured",
        )


def _supabase_headers(
    *,
    content_type: str | None = None,
    prefer: str | None = None,
    include_profiles: bool = False,
) -> dict[str, str]:
    _assert_config()
    assert SUPABASE_SERVICE_ROLE_KEY is not None

    headers = {"apikey": SUPABASE_SERVICE_ROLE_KEY}
    if not _is_supabase_opaque_key(SUPABASE_SERVICE_ROLE_KEY):
        headers["Authorization"] = f"Bearer {SUPABASE_SERVICE_ROLE_KEY}"
    if content_type:
        headers["Content-Type"] = content_type
    if include_profiles:
        headers["Accept-Profile"] = SUPABASE_DB_SCHEMA
        headers["Content-Profile"] = SUPABASE_DB_SCHEMA
    if prefer:
        headers["Prefer"] = prefer
    return headers


def _validate_content_type(resume_path: str, content_type_header: str | None) -> None:
    if not content_type_header:
        return

    content_type = content_type_header.split(";", 1)[0].strip().lower()
    lower_path = resume_path.lower()

    if lower_path.endswith(".pdf"):
        allowed = ALLOWED_PDF_CONTENT_TYPES
    elif lower_path.endswith(".docx"):
        allowed = ALLOWED_DOCX_CONTENT_TYPES
    else:
        return

    if content_type not in allowed:
        raise HttpError(
            status=415,
            code=ERR_UNSUPPORTED_CONTENT_TYPE,
            message=f"Unexpected content-type '{content_type}' for file '{resume_path}'",
        )


def _validate_file_size(content_length_header: str | None, file_bytes_len: int) -> None:
    if content_length_header:
        try:
            content_length = int(content_length_header)
            if content_length > MAX_RESUME_FILE_SIZE_BYTES:
                raise HttpError(
                    status=413,
                    code=ERR_FILE_TOO_LARGE,
                    message=f"Resume file exceeds max size of {MAX_RESUME_FILE_SIZE_MB}MB",
                )
        except ValueError:
            pass

    if file_bytes_len > MAX_RESUME_FILE_SIZE_BYTES:
        raise HttpError(
            status=413,
            code=ERR_FILE_TOO_LARGE,
            message=f"Resume file exceeds max size of {MAX_RESUME_FILE_SIZE_MB}MB",
        )


async def _supabase_fetch(path: str, method: str, body: dict[str, Any] | None = None, prefer: str | None = None) -> Any:
    _assert_config()
    headers = _supabase_headers(content_type="application/json", prefer=prefer, include_profiles=True)
    url = f"{SUPABASE_URL}/rest/v1/{path}"
    async with httpx.AsyncClient(timeout=30.0) as client:
        response = await client.request(method=method, url=url, headers=headers, json=body)

    if response.status_code >= 400:
        raise HttpError(
            status=502,
            code=ERR_DB_ERROR,
            message=f"Supabase error: {response.text}",
        )

    if "application/json" in response.headers.get("content-type", ""):
        return response.json()
    return None


async def _supabase_count(path: str) -> int:
    _assert_config()
    headers = _supabase_headers(prefer="count=exact", include_profiles=True)
    url = f"{SUPABASE_URL}/rest/v1/{path}"
    async with httpx.AsyncClient(timeout=30.0) as client:
        response = await client.request(method="HEAD", url=url, headers=headers)

    if response.status_code >= 400:
        raise HttpError(
            status=502,
            code=ERR_DB_ERROR,
            message=f"Supabase count error: {response.text}",
        )

    content_range = response.headers.get("content-range", "")
    if "/" not in content_range:
        raise HttpError(
            status=500,
            code=ERR_DB_ERROR,
            message="Supabase count response missing content-range header",
        )
    try:
        return int(content_range.rsplit("/", 1)[1])
    except ValueError as exc:
        raise HttpError(
            status=500,
            code=ERR_DB_ERROR,
            message="Supabase count response had invalid content-range header",
        ) from exc


async def _register_worker_presence() -> None:
    await _supabase_fetch(
        path="extract_worker_heartbeats?on_conflict=worker_id",
        method="POST",
        body={
            "worker_id": WORKER_CLAIM_OWNER,
            "hostname": socket.gethostname(),
            "pid": os.getpid(),
            "role": "extractor",
            "started_at": _utc_now_iso(),
            "last_seen_at": _utc_now_iso(),
        },
        prefer="resolution=merge-duplicates,return=minimal",
    )


async def _heartbeat_worker_presence() -> None:
    path = f"extract_worker_heartbeats?worker_id=eq.{_encode_rest_eq_value(WORKER_CLAIM_OWNER)}"
    rows = await _supabase_fetch(
        path=path,
        method="PATCH",
        body={"last_seen_at": _utc_now_iso()},
        prefer="return=representation",
    )
    if not isinstance(rows, list) or len(rows) == 0:
        await _register_worker_presence()


async def _release_worker_presence() -> None:
    path = f"extract_worker_heartbeats?worker_id=eq.{_encode_rest_eq_value(WORKER_CLAIM_OWNER)}"
    await _supabase_fetch(path=path, method="DELETE", prefer="return=minimal")


async def _run_worker_presence_heartbeat() -> None:
    while True:
        await asyncio.sleep(WORKER_HEARTBEAT_INTERVAL_SECONDS)
        await _heartbeat_worker_presence()


async def _reset_stale_runs(limit: int | None = None) -> list[dict[str, Any]]:
    rows = await _supabase_fetch(
        path="rpc/reset_stale_resume_runs",
        method="POST",
        body={
            "p_stale_seconds": WORKER_LEASE_SECONDS,
            "p_limit": limit if limit is not None else WORKER_STALE_RESET_LIMIT,
        },
    )
    if rows is None:
        return []
    if not isinstance(rows, list):
        raise HttpError(
            status=500,
            code=ERR_INVALID_RUN_ROW,
            message="Reset-stale RPC returned invalid payload",
        )
    return rows


async def _run_stale_reset_if_due(last_reset_at: float) -> float:
    now = time.monotonic()
    if now - last_reset_at < WORKER_STALE_RESET_INTERVAL_SECONDS:
        return last_reset_at
    rows = await _reset_stale_runs()
    if rows:
        _log_event(phase="worker_stale_reset", reset_count=len(rows))
    return now


async def _claim_run(run_id: UUID) -> dict[str, Any]:
    rows = await _supabase_fetch(
        path="rpc/claim_resume_run",
        method="POST",
        body={
            "p_run_id": str(run_id),
            "p_claimed_by": API_CLAIM_OWNER,
        },
    )
    if not isinstance(rows, list) or len(rows) == 0:
        raise HttpError(
            status=409,
            code=ERR_RUN_NOT_QUEUED,
            message="Run is not in queued state or does not exist",
        )
    return rows[0]


async def _claim_next_run() -> dict[str, Any] | None:
    rows = await _supabase_fetch(
        path="rpc/claim_next_resume_run",
        method="POST",
        body={
            "p_claimed_by": WORKER_CLAIM_OWNER,
            "p_lease_seconds": WORKER_LEASE_SECONDS,
        },
    )
    if rows is None:
        return None
    if not isinstance(rows, list):
        raise HttpError(
            status=500,
            code=ERR_INVALID_RUN_ROW,
            message="Claim-next RPC returned invalid payload",
        )
    if len(rows) == 0:
        return None
    return rows[0]


async def _insert_resume_document(
    run_id: UUID,
    run_user_id: str,
    run_resume_path: str,
    extracted_text: str,
    metadata: dict[str, Any],
) -> None:
    # Idempotent write for retry flows: if a row already exists for run_id,
    # overwrite it for the same run instead of failing on unique constraint.
    await _supabase_fetch(
        path="resume_documents?on_conflict=run_id",
        method="POST",
        body={
            "run_id": str(run_id),
            "user_id": run_user_id,
            "resume_path": run_resume_path,
            "text": extracted_text,
            "text_source": "extract_service",
            "metadata": metadata,
        },
        prefer="resolution=merge-duplicates,return=minimal",
    )


def _encode_rest_eq_value(value: str) -> str:
    return quote(value, safe="-_.~:")


def _utc_now_iso() -> str:
    return datetime.now(timezone.utc).isoformat()


def _parse_iso_datetime(value: str | None) -> datetime | None:
    if not value or not isinstance(value, str):
        return None
    try:
        parsed = datetime.fromisoformat(value.replace("Z", "+00:00"))
    except ValueError:
        return None
    if parsed.tzinfo is None:
        return parsed.replace(tzinfo=timezone.utc)
    return parsed


def _compute_queue_wait_ms(created_at_raw: str | None) -> int | None:
    created_at = _parse_iso_datetime(created_at_raw)
    if created_at is None:
        return None
    return max(0, int((datetime.now(timezone.utc) - created_at).total_seconds() * 1000))


def _failure_bucket(error_code: str) -> str:
    if error_code in {ERR_DB_ERROR}:
        return "db"
    if error_code in {ERR_DOWNLOAD_FAILED}:
        return "download"
    if error_code in {ERR_PARSE_ERROR, ERR_EMPTY_EXTRACTED_TEXT, ERR_UNSUPPORTED_FILE_TYPE, ERR_UNSUPPORTED_CONTENT_TYPE}:
        return "parse"
    if error_code in {ERR_FILE_TOO_LARGE, ERR_EMPTY_RESUME_FILE}:
        return "input"
    if error_code in {ERR_OCR_UNAVAILABLE, ERR_OCR_DEPENDENCY_MISSING}:
        return "ocr"
    if error_code in {ERR_MISSING_CONFIG, ERR_INVALID_RUN_ROW, ERR_RUN_NOT_QUEUED, ERR_RUN_NOT_EXTRACTING}:
        return "system"
    return "other"


def _log_extraction_summary(
    *,
    run_id: UUID,
    status: str,
    duration_ms: int,
    parser: str | None = None,
    attempt_count: int,
    queue_wait_ms: int | None,
    error_code: str | None = None,
    failure_bucket: str | None = None,
) -> None:
    _log_event(
        phase="extract_summary",
        run_id=run_id,
        parser=parser,
        duration_ms=duration_ms,
        status=status,
        attempt_count=attempt_count,
        queue_wait_ms=queue_wait_ms,
        error_code=error_code,
        failure_bucket=failure_bucket,
    )


def _log_generation_summary(
    *,
    run_id: UUID,
    status: str,
    duration_ms: int,
    attempt_count: int,
    queue_wait_ms: int | None,
    error_code: str | None = None,
) -> None:
    _log_event(
        phase="generate_summary",
        run_id=run_id,
        duration_ms=duration_ms,
        status=status,
        attempt_count=attempt_count,
        queue_wait_ms=queue_wait_ms,
        error_code=error_code,
    )


def _log_pdf_summary(
    *,
    run_id: UUID,
    status: str,
    duration_ms: int,
    attempt_count: int,
    queue_wait_ms: int | None,
    error_code: str | None = None,
) -> None:
    _log_event(
        phase="pdf_summary",
        run_id=run_id,
        duration_ms=duration_ms,
        status=status,
        attempt_count=attempt_count,
        queue_wait_ms=queue_wait_ms,
        error_code=error_code,
    )


def _resume_runs_count_path(*filters: str) -> str:
    query = "&".join(["select=id", *filters])
    return f"resume_runs?{query}"


async def _get_queue_snapshot() -> dict[str, int]:
    now = datetime.now(timezone.utc)
    cutoff_1m = (now - timedelta(minutes=1)).isoformat()
    cutoff_5m = (now - timedelta(minutes=5)).isoformat()
    cutoff_15m = (now - timedelta(minutes=15)).isoformat()
    active_worker_cutoff = (now - timedelta(seconds=max(WORKER_ACTIVE_WINDOW_SECONDS, 1))).isoformat()
    stale_extracting_cutoff = (now - timedelta(seconds=max(WORKER_LEASE_SECONDS, 1))).isoformat()
    stale_generating_cutoff = (now - timedelta(seconds=max(GENERATION_WORKER_LEASE_SECONDS, 1))).isoformat()
    stale_pdf_cutoff = (now - timedelta(seconds=max(PDF_WORKER_LEASE_SECONDS, 1))).isoformat()
    queued_total = await _supabase_count(_resume_runs_count_path("status=eq.queued"))
    queued_over_1m = await _supabase_count(
        _resume_runs_count_path("status=eq.queued", f"created_at=lte.{quote(cutoff_1m, safe='')}")
    )
    queued_over_5m = await _supabase_count(
        _resume_runs_count_path("status=eq.queued", f"created_at=lte.{quote(cutoff_5m, safe='')}")
    )
    queued_over_15m = await _supabase_count(
        _resume_runs_count_path("status=eq.queued", f"created_at=lte.{quote(cutoff_15m, safe='')}")
    )
    extracting_total = await _supabase_count(_resume_runs_count_path("status=eq.extracting"))
    stale_extracting_total = await _supabase_count(
        _resume_runs_count_path(
            "status=eq.extracting",
            f"extraction_heartbeat_at=lte.{quote(stale_extracting_cutoff, safe='')}",
        )
    )
    queued_generate_total = await _supabase_count(_resume_runs_count_path("status=eq.queued_generate"))
    queued_generate_over_1m = await _supabase_count(
        _resume_runs_count_path("status=eq.queued_generate", f"updated_at=lte.{quote(cutoff_1m, safe='')}")
    )
    queued_generate_over_5m = await _supabase_count(
        _resume_runs_count_path("status=eq.queued_generate", f"updated_at=lte.{quote(cutoff_5m, safe='')}")
    )
    generating_total = await _supabase_count(_resume_runs_count_path("status=eq.generating"))
    stale_generating_total = await _supabase_count(
        _resume_runs_count_path(
            "status=eq.generating",
            f"generation_heartbeat_at=lte.{quote(stale_generating_cutoff, safe='')}",
        )
    )
    queued_pdf_total = await _supabase_count(_resume_runs_count_path("status=eq.queued_pdf"))
    queued_pdf_over_1m = await _supabase_count(
        _resume_runs_count_path("status=eq.queued_pdf", f"updated_at=lte.{quote(cutoff_1m, safe='')}")
    )
    queued_pdf_over_5m = await _supabase_count(
        _resume_runs_count_path("status=eq.queued_pdf", f"updated_at=lte.{quote(cutoff_5m, safe='')}")
    )
    compiling_pdf_total = await _supabase_count(_resume_runs_count_path("status=eq.compiling_pdf"))
    stale_compiling_pdf_total = await _supabase_count(
        _resume_runs_count_path(
            "status=eq.compiling_pdf",
            f"pdf_heartbeat_at=lte.{quote(stale_pdf_cutoff, safe='')}",
        )
    )
    completed_total = await _supabase_count(_resume_runs_count_path("status=eq.completed"))
    failed_total = await _supabase_count(_resume_runs_count_path("status=eq.failed"))
    active_workers = await _supabase_count(
        "extract_worker_heartbeats?"
        + "&".join(
            [
                "select=worker_id",
                f"last_seen_at=gte.{quote(active_worker_cutoff, safe='')}",
            ]
        )
    )
    return {
        "queued_total": queued_total,
        "queued_over_1m": queued_over_1m,
        "queued_over_5m": queued_over_5m,
        "queued_over_15m": queued_over_15m,
        "extracting_total": extracting_total,
        "stale_extracting_total": stale_extracting_total,
        "queued_generate_total": queued_generate_total,
        "queued_generate_over_1m": queued_generate_over_1m,
        "queued_generate_over_5m": queued_generate_over_5m,
        "generating_total": generating_total,
        "stale_generating_total": stale_generating_total,
        "queued_pdf_total": queued_pdf_total,
        "queued_pdf_over_1m": queued_pdf_over_1m,
        "queued_pdf_over_5m": queued_pdf_over_5m,
        "compiling_pdf_total": compiling_pdf_total,
        "stale_compiling_pdf_total": stale_compiling_pdf_total,
        "completed_total": completed_total,
        "failed_total": failed_total,
        "active_workers": active_workers,
    }


async def _set_run_ready_for_generation(run_id: UUID, claimed_by: str) -> None:
    path = (
        f"resume_runs?id=eq.{run_id}&status=eq.extracting"
        f"&extraction_claimed_by=eq.{_encode_rest_eq_value(claimed_by)}"
    )
    rows = await _supabase_fetch(
        path=path,
        method="PATCH",
        body={
            "status": "queued_generate",
            "error_code": None,
            "error_message": None,
            "extraction_claimed_by": None,
            "extraction_claimed_at": None,
            "extraction_heartbeat_at": None,
        },
        prefer="return=representation",
    )
    if not isinstance(rows, list) or len(rows) == 0:
        raise HttpError(
            status=409,
            code=ERR_RUN_NOT_EXTRACTING,
            message="Run is not in extracting state during finalize",
        )


async def _set_run_failed(run_id: UUID, claimed_by: str, code: str, message: str) -> None:
    path = (
        f"resume_runs?id=eq.{run_id}&status=eq.extracting"
        f"&extraction_claimed_by=eq.{_encode_rest_eq_value(claimed_by)}"
    )
    rows = await _supabase_fetch(
        path=path,
        method="PATCH",
        body={
            "status": "failed",
            "error_code": code,
            "error_message": message[:400],
            "extraction_claimed_by": None,
            "extraction_claimed_at": None,
            "extraction_heartbeat_at": None,
            "extraction_next_retry_at": None,
        },
        prefer="return=representation",
    )
    if not isinstance(rows, list) or len(rows) == 0:
        raise HttpError(
            status=409,
            code=ERR_RUN_NOT_EXTRACTING,
            message="Run is not in extracting state during failure finalize",
        )


async def _requeue_run(run_id: UUID, claimed_by: str, code: str, message: str, delay_seconds: float) -> None:
    retry_at = datetime.now(timezone.utc).timestamp() + max(delay_seconds, 1)
    path = (
        f"resume_runs?id=eq.{run_id}&status=eq.extracting"
        f"&extraction_claimed_by=eq.{_encode_rest_eq_value(claimed_by)}"
    )
    rows = await _supabase_fetch(
        path=path,
        method="PATCH",
        body={
            "status": "queued",
            "error_code": code,
            "error_message": message[:400],
            "extraction_claimed_by": None,
            "extraction_claimed_at": None,
            "extraction_heartbeat_at": None,
            "extraction_next_retry_at": datetime.fromtimestamp(retry_at, timezone.utc).isoformat(),
        },
        prefer="return=representation",
    )
    if not isinstance(rows, list) or len(rows) == 0:
        raise HttpError(
            status=409,
            code=ERR_RUN_NOT_EXTRACTING,
            message="Run is not in extracting state during retry requeue",
        )


def _is_retryable_error(error: Exception, error_code: str) -> bool:
    if isinstance(error, HttpError):
        if error_code in {ERR_DB_ERROR, ERR_DOWNLOAD_FAILED}:
            return True
        return error.status >= 500 and error_code not in {ERR_MISSING_CONFIG, ERR_OCR_DEPENDENCY_MISSING}
    return False


async def _heartbeat_claimed_run(run_id: UUID, claimed_by: str) -> None:
    path = (
        f"resume_runs?id=eq.{run_id}&status=eq.extracting"
        f"&extraction_claimed_by=eq.{_encode_rest_eq_value(claimed_by)}"
    )
    rows = await _supabase_fetch(
        path=path,
        method="PATCH",
        body={"extraction_heartbeat_at": _utc_now_iso()},
        prefer="return=representation",
    )
    if not isinstance(rows, list) or len(rows) == 0:
        raise HttpError(
            status=409,
            code=ERR_RUN_NOT_EXTRACTING,
            message="Run is not in extracting state during heartbeat",
        )


async def _run_claim_heartbeat(run_id: UUID, claimed_by: str) -> None:
    while True:
        await asyncio.sleep(RUN_HEARTBEAT_INTERVAL_SECONDS)
        await _heartbeat_claimed_run(run_id, claimed_by)


async def _claim_next_generate_run() -> dict[str, Any] | None:
    rows = await _supabase_fetch(
        path="rpc/claim_next_generate_run",
        method="POST",
        body={"p_claimed_by": GENERATION_WORKER_CLAIM_OWNER, "p_lease_seconds": GENERATION_WORKER_LEASE_SECONDS},
    )
    if rows is None:
        return None
    if not isinstance(rows, list):
        raise HttpError(status=500, code=ERR_INVALID_RUN_ROW, message="claim_next_generate_run returned invalid payload")
    if not rows:
        return None
    row = rows[0]
    if not isinstance(row, dict):
        raise HttpError(status=500, code=ERR_INVALID_RUN_ROW, message="claim_next_generate_run returned invalid row")
    return row


async def _heartbeat_claimed_generate_run(run_id: UUID, claimed_by: str) -> None:
    path = (
        f"resume_runs?id=eq.{run_id}&status=eq.generating"
        f"&generation_claimed_by=eq.{_encode_rest_eq_value(claimed_by)}"
    )
    rows = await _supabase_fetch(
        path=path,
        method="PATCH",
        body={"generation_heartbeat_at": _utc_now_iso()},
        prefer="return=representation",
    )
    if not isinstance(rows, list) or len(rows) == 0:
        raise HttpError(
            status=409,
            code=ERR_GENERATE_NOT_CLAIMED,
            message="Run is not in generating state during generation heartbeat",
        )


async def _run_generation_claim_heartbeat(run_id: UUID, claimed_by: str) -> None:
    while True:
        await asyncio.sleep(GENERATION_RUN_HEARTBEAT_INTERVAL_SECONDS)
        await _heartbeat_claimed_generate_run(run_id, claimed_by)


async def _requeue_generate_run(run_id: UUID, claimed_by: str, code: str, message: str, delay_seconds: float) -> None:
    retry_at = datetime.now(timezone.utc).timestamp() + max(delay_seconds, 1)
    path = (
        f"resume_runs?id=eq.{run_id}&status=eq.generating"
        f"&generation_claimed_by=eq.{_encode_rest_eq_value(claimed_by)}"
    )
    rows = await _supabase_fetch(
        path=path,
        method="PATCH",
        body={
            "status": "queued_generate",
            "error_code": code,
            "error_message": message[:400],
            "generation_claimed_by": None,
            "generation_claimed_at": None,
            "generation_heartbeat_at": None,
            "generation_next_retry_at": datetime.fromtimestamp(retry_at, timezone.utc).isoformat(),
        },
        prefer="return=representation",
    )
    if not isinstance(rows, list) or len(rows) == 0:
        raise HttpError(
            status=409,
            code=ERR_GENERATE_NOT_CLAIMED,
            message="Run is not in generating state during generation retry requeue",
        )


async def _set_generate_failed(run_id: UUID, claimed_by: str, code: str, message: str) -> None:
    path = (
        f"resume_runs?id=eq.{run_id}&status=eq.generating"
        f"&generation_claimed_by=eq.{_encode_rest_eq_value(claimed_by)}"
    )
    rows = await _supabase_fetch(
        path=path,
        method="PATCH",
        body={
            "status": "failed",
            "error_code": code,
            "error_message": message[:400],
            "generation_claimed_by": None,
            "generation_claimed_at": None,
            "generation_heartbeat_at": None,
            "generation_next_retry_at": None,
        },
        prefer="return=representation",
    )
    if not isinstance(rows, list) or len(rows) == 0:
        raise HttpError(
            status=409,
            code=ERR_GENERATE_NOT_CLAIMED,
            message="Run is not in generating state during generation failure finalize",
        )


async def _invoke_supabase_function(function_name: str, body: dict[str, Any]) -> dict[str, Any]:
    _assert_config()
    url = f"{SUPABASE_URL}/functions/v1/{function_name}"
    headers = _supabase_headers(content_type="application/json")
    async with httpx.AsyncClient(timeout=120.0) as client:
        response = await client.post(url, headers=headers, json=body)

    payload = response.json() if "application/json" in response.headers.get("content-type", "") else None
    if response.status_code >= 400:
        error_message = None
        if isinstance(payload, dict):
          error_message = payload.get("error_message")
        raise HttpError(
            status=502,
            code=ERR_FUNCTION_CALL_FAILED,
            message=f"{function_name} failed: {error_message or response.text}",
        )
    if not isinstance(payload, dict):
        raise HttpError(status=502, code=ERR_FUNCTION_CALL_FAILED, message=f"{function_name} returned invalid payload")
    return payload


async def _process_claimed_generate_run(
    run_id: UUID,
    request_id: str,
    claimed_by: str,
    attempt_count: int,
    queue_wait_ms: int | None,
) -> None:
    request_start = time.perf_counter()
    heartbeat_task = asyncio.create_task(_run_generation_claim_heartbeat(run_id, claimed_by))
    try:
        bullets_start = time.perf_counter()
        await _invoke_supabase_function(
            "generate-bullets",
            {"run_id": str(run_id), "request_id": request_id},
        )
        _log_event(
            phase="generate_bullets_complete",
            run_id=run_id,
            duration_ms=int((time.perf_counter() - bullets_start) * 1000),
        )

        tailored_start = time.perf_counter()
        await _invoke_supabase_function(
            "generate-tailored-resume",
            {"run_id": str(run_id), "request_id": request_id},
        )
        _log_event(
            phase="generate_tailored_resume_complete",
            run_id=run_id,
            duration_ms=int((time.perf_counter() - tailored_start) * 1000),
        )
        _log_event(
            phase="generate_success",
            run_id=run_id,
            duration_ms=int((time.perf_counter() - request_start) * 1000),
            attempt_count=attempt_count,
            queue_wait_ms=queue_wait_ms,
        )
        _log_generation_summary(
            run_id=run_id,
            status="success",
            duration_ms=int((time.perf_counter() - request_start) * 1000),
            attempt_count=attempt_count,
            queue_wait_ms=queue_wait_ms,
        )
    except Exception as error:
        error_code = error.code if isinstance(error, HttpError) else ERR_FUNCTION_CALL_FAILED
        error_message = str(error) if str(error) else "Unknown generation failure"
        _log_event(
            phase="generate_error",
            run_id=run_id,
            duration_ms=int((time.perf_counter() - request_start) * 1000),
            error_code=error_code,
            error_message=error_message,
            attempt_count=attempt_count,
            queue_wait_ms=queue_wait_ms,
        )
        if attempt_count < GENERATION_WORKER_MAX_ATTEMPTS:
            await _requeue_generate_run(
                run_id,
                claimed_by,
                error_code,
                error_message,
                GENERATION_RETRY_DELAY_SECONDS,
            )
        else:
            await _set_generate_failed(run_id, claimed_by, error_code, error_message)
        _log_generation_summary(
            run_id=run_id,
            status="failed" if attempt_count >= GENERATION_WORKER_MAX_ATTEMPTS else "requeued",
            duration_ms=int((time.perf_counter() - request_start) * 1000),
            attempt_count=attempt_count,
            queue_wait_ms=queue_wait_ms,
            error_code=error_code,
        )
        raise
    finally:
        heartbeat_task.cancel()
        with suppress(asyncio.CancelledError):
            await heartbeat_task


async def _claim_next_pdf_run() -> dict[str, Any] | None:
    rows = await _supabase_fetch(
        path="rpc/claim_next_pdf_run",
        method="POST",
        body={"p_claimed_by": PDF_WORKER_CLAIM_OWNER, "p_lease_seconds": PDF_WORKER_LEASE_SECONDS},
    )
    if rows is None:
        return None
    if not isinstance(rows, list):
        raise HttpError(status=500, code=ERR_INVALID_RUN_ROW, message="claim_next_pdf_run returned invalid payload")
    if not rows:
        return None
    row = rows[0]
    if not isinstance(row, dict):
        raise HttpError(status=500, code=ERR_INVALID_RUN_ROW, message="claim_next_pdf_run returned invalid row")
    return row


async def _heartbeat_claimed_pdf_run(run_id: UUID, claimed_by: str) -> None:
    path = (
        f"resume_runs?id=eq.{run_id}&status=eq.compiling_pdf"
        f"&pdf_claimed_by=eq.{_encode_rest_eq_value(claimed_by)}"
    )
    rows = await _supabase_fetch(
        path=path,
        method="PATCH",
        body={"pdf_heartbeat_at": _utc_now_iso()},
        prefer="return=representation",
    )
    if not isinstance(rows, list) or len(rows) == 0:
        raise HttpError(
            status=409,
            code=ERR_PDF_NOT_CLAIMED,
            message="Run is not in compiling_pdf state during pdf heartbeat",
        )


async def _run_pdf_claim_heartbeat(run_id: UUID, claimed_by: str) -> None:
    while True:
        await asyncio.sleep(PDF_RUN_HEARTBEAT_INTERVAL_SECONDS)
        await _heartbeat_claimed_pdf_run(run_id, claimed_by)


async def _requeue_pdf_run(run_id: UUID, claimed_by: str, code: str, message: str, delay_seconds: float) -> None:
    retry_at = datetime.now(timezone.utc).timestamp() + max(delay_seconds, 1)
    path = (
        f"resume_runs?id=eq.{run_id}&status=eq.compiling_pdf"
        f"&pdf_claimed_by=eq.{_encode_rest_eq_value(claimed_by)}"
    )
    rows = await _supabase_fetch(
        path=path,
        method="PATCH",
        body={
            "status": "queued_pdf",
            "error_code": code,
            "error_message": message[:400],
            "pdf_claimed_by": None,
            "pdf_claimed_at": None,
            "pdf_heartbeat_at": None,
            "pdf_next_retry_at": datetime.fromtimestamp(retry_at, timezone.utc).isoformat(),
        },
        prefer="return=representation",
    )
    if not isinstance(rows, list) or len(rows) == 0:
        raise HttpError(
            status=409,
            code=ERR_PDF_NOT_CLAIMED,
            message="Run is not in compiling_pdf state during pdf retry requeue",
        )


async def _set_pdf_failed(run_id: UUID, claimed_by: str, code: str, message: str) -> None:
    path = (
        f"resume_runs?id=eq.{run_id}&status=eq.compiling_pdf"
        f"&pdf_claimed_by=eq.{_encode_rest_eq_value(claimed_by)}"
    )
    rows = await _supabase_fetch(
        path=path,
        method="PATCH",
        body={
            "status": "failed",
            "error_code": code,
            "error_message": message[:400],
            "pdf_claimed_by": None,
            "pdf_claimed_at": None,
            "pdf_heartbeat_at": None,
            "pdf_next_retry_at": None,
        },
        prefer="return=representation",
    )
    if not isinstance(rows, list) or len(rows) == 0:
        raise HttpError(
            status=409,
            code=ERR_PDF_NOT_CLAIMED,
            message="Run is not in compiling_pdf state during pdf failure finalize",
        )


async def _process_claimed_pdf_run(
    run_id: UUID,
    attempt_count: int,
    queue_wait_ms: int | None,
) -> None:
    request_start = time.perf_counter()
    heartbeat_task = asyncio.create_task(_run_pdf_claim_heartbeat(run_id, PDF_WORKER_CLAIM_OWNER))
    try:
        compile_start = time.perf_counter()
        await _invoke_supabase_function(
            "compile-tailored-resume-pdf",
            {"run_id": str(run_id)},
        )
        _log_event(
            phase="pdf_compile_complete",
            run_id=run_id,
            duration_ms=int((time.perf_counter() - compile_start) * 1000),
        )
        _log_event(
            phase="pdf_success",
            run_id=run_id,
            duration_ms=int((time.perf_counter() - request_start) * 1000),
            attempt_count=attempt_count,
            queue_wait_ms=queue_wait_ms,
        )
        _log_pdf_summary(
            run_id=run_id,
            status="success",
            duration_ms=int((time.perf_counter() - request_start) * 1000),
            attempt_count=attempt_count,
            queue_wait_ms=queue_wait_ms,
        )
    except Exception as error:
        error_code = error.code if isinstance(error, HttpError) else ERR_FUNCTION_CALL_FAILED
        error_message = str(error) if str(error) else "Unknown pdf generation failure"
        _log_event(
            phase="pdf_error",
            run_id=run_id,
            duration_ms=int((time.perf_counter() - request_start) * 1000),
            error_code=error_code,
            error_message=error_message,
            attempt_count=attempt_count,
            queue_wait_ms=queue_wait_ms,
        )
        if attempt_count < PDF_WORKER_MAX_ATTEMPTS:
            await _requeue_pdf_run(
                run_id,
                PDF_WORKER_CLAIM_OWNER,
                error_code,
                error_message,
                PDF_RETRY_DELAY_SECONDS,
            )
        else:
            await _set_pdf_failed(run_id, PDF_WORKER_CLAIM_OWNER, error_code, error_message)
        _log_pdf_summary(
            run_id=run_id,
            status="failed" if attempt_count >= PDF_WORKER_MAX_ATTEMPTS else "requeued",
            duration_ms=int((time.perf_counter() - request_start) * 1000),
            attempt_count=attempt_count,
            queue_wait_ms=queue_wait_ms,
            error_code=error_code,
        )
        raise
    finally:
        heartbeat_task.cancel()
        with suppress(asyncio.CancelledError):
            await heartbeat_task


def _normalize_line(line: str) -> str:
    line = line.replace("\x00", " ").replace("\xa0", " ")
    line = ZERO_WIDTH_CHARS_RE.sub("", line)
    line = line.replace("\t", " ")
    line = re.sub(r"\s+", " ", line).strip()
    if not line:
        return ""
    if PAGE_MARKER_RE.match(line):
        return ""

    # Normalize bullet styles into a single durable format for chunking.
    bullet_removed = BULLET_PREFIX_RE.sub("", line).strip()
    if bullet_removed and bullet_removed != line:
        return f"- {bullet_removed}"
    return line


def _normalize_text(raw_text: str, remove_repeated_boilerplate: bool = False) -> str:
    lines = [_normalize_line(line) for line in raw_text.splitlines()]
    lines = [line for line in lines if line]

    if remove_repeated_boilerplate and lines:
        lowered = [line.casefold() for line in lines]
        freq = Counter(lowered)
        kept_boilerplate: set[str] = set()
        filtered: list[str] = []
        for line in lines:
            key = line.casefold()
            is_repeated_boilerplate = (
                freq[key] >= 3
                and len(line) <= 120
                and len(line.split()) <= 12
            )
            if is_repeated_boilerplate:
                if key in kept_boilerplate:
                    continue
                kept_boilerplate.add(key)
            filtered.append(line)
        lines = filtered

    # Remove consecutive duplicate lines to reduce OCR and parser noise.
    deduped: list[str] = []
    for line in lines:
        if deduped and deduped[-1].casefold() == line.casefold():
            continue
        deduped.append(line)

    return "\n".join(deduped).strip()


async def _download_resume_bytes(resume_path: str) -> bytes:
    _assert_config()
    encoded_path = quote(resume_path, safe="/")
    url = f"{SUPABASE_URL}/storage/v1/object/{SUPABASE_STORAGE_BUCKET}/{encoded_path}"
    headers = _supabase_headers()
    try:
        async with httpx.AsyncClient(timeout=EXTRACT_DOWNLOAD_TIMEOUT_SECONDS) as client:
            response = await client.get(url, headers=headers)
    except httpx.TimeoutException as exc:
        raise HttpError(
            status=504,
            code=ERR_DOWNLOAD_FAILED,
            message="Timed out while downloading resume file from storage",
        ) from exc
    except httpx.HTTPError as exc:
        raise HttpError(
            status=502,
            code=ERR_DOWNLOAD_FAILED,
            message=f"Network error while downloading resume file: {exc}",
        ) from exc

    if response.status_code >= 400:
        raise HttpError(
            status=502,
            code=ERR_DOWNLOAD_FAILED,
            message=f"Failed to download resume file from storage: {response.text}",
        )

    _validate_content_type(resume_path, response.headers.get("content-type"))
    _validate_file_size(response.headers.get("content-length"), len(response.content))

    if not response.content:
        raise HttpError(
            status=422,
            code=ERR_EMPTY_RESUME_FILE,
            message="Resume file is empty",
        )
    return response.content


def _extract_pdf(file_bytes: bytes) -> tuple[str, dict[str, Any]]:
    try:
        reader = PdfReader(BytesIO(file_bytes))
    except Exception as exc:
        raise HttpError(
            status=422,
            code=ERR_PARSE_ERROR,
            message=f"Failed to parse PDF: {exc}",
        ) from exc
    pages_text: list[str] = []
    for page in reader.pages:
        try:
            pages_text.append(page.extract_text() or "")
        except Exception:
            pages_text.append("")
    merged = "\n".join(pages_text)
    cleaned = _normalize_text(merged, remove_repeated_boilerplate=True)
    if cleaned:
        return cleaned, {"parser": "pypdf", "pages": len(reader.pages), "ocr_used": False}

    return _extract_pdf_with_ocr(file_bytes, len(reader.pages))


def _extract_pdf_with_ocr(file_bytes: bytes, page_count: int) -> tuple[str, dict[str, Any]]:
    if shutil.which("tesseract") is None:
        raise HttpError(
            status=422,
            code=ERR_OCR_UNAVAILABLE,
            message="No extractable text found in PDF and OCR is unavailable (missing tesseract binary)",
        )
    try:
        # Lazy imports so normal text-PDF flow does not require OCR packages at import time.
        from pdf2image import convert_from_bytes
        import pytesseract
    except ImportError as exc:
        raise HttpError(
            status=500,
            code=ERR_OCR_DEPENDENCY_MISSING,
            message="OCR dependencies are not installed (pdf2image and pytesseract)",
        ) from exc

    try:
        images = convert_from_bytes(file_bytes, fmt="png")
    except Exception as exc:
        raise HttpError(
            status=422,
            code=ERR_PARSE_ERROR,
            message=f"Failed to render PDF pages for OCR: {exc}",
        ) from exc

    ocr_pages: list[str] = []
    for image in images:
        try:
            ocr_pages.append(pytesseract.image_to_string(image))
        except Exception:
            ocr_pages.append("")

    cleaned = _normalize_text("\n".join(ocr_pages), remove_repeated_boilerplate=True)
    if not cleaned:
        raise HttpError(
            status=422,
            code=ERR_EMPTY_EXTRACTED_TEXT,
            message="No extractable text found in PDF after OCR",
        )
    return cleaned, {"parser": "pypdf+tesseract", "pages": page_count, "ocr_used": True}


def _extract_docx(file_bytes: bytes) -> tuple[str, dict[str, Any]]:
    try:
        document = Document(BytesIO(file_bytes))
    except Exception as exc:
        raise HttpError(
            status=422,
            code=ERR_PARSE_ERROR,
            message=f"Failed to parse DOCX: {exc}",
        ) from exc
    chunks: list[str] = []

    for paragraph in document.paragraphs:
        if paragraph.text and paragraph.text.strip():
            chunks.append(paragraph.text)

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text and cell.text.strip()]
            if cells:
                chunks.append(" | ".join(cells))

    cleaned = _normalize_text("\n".join(chunks))
    if not cleaned:
        raise HttpError(
            status=422,
            code=ERR_EMPTY_EXTRACTED_TEXT,
            message="No extractable text found in DOCX",
        )
    return cleaned, {"parser": "python-docx", "paragraph_count": len(document.paragraphs)}


async def _perform_extraction(resume_path: str) -> tuple[str, dict[str, Any]]:
    file_bytes = await _download_resume_bytes(resume_path)
    lower_path = resume_path.lower()

    if lower_path.endswith(".pdf"):
        return _extract_pdf(file_bytes)
    if lower_path.endswith(".docx"):
        return _extract_docx(file_bytes)

    raise HttpError(
        status=422,
        code=ERR_UNSUPPORTED_FILE_TYPE,
        message="Only PDF and DOCX files are supported",
    )


async def _process_claimed_run(
    run_id: UUID,
    run_user_id: str,
    run_resume_path: str,
    claimed_by: str,
    attempt_count: int,
    queue_wait_ms: int | None,
) -> None:
    parser_name: str | None = None
    request_start = time.perf_counter()
    heartbeat_task = asyncio.create_task(_run_claim_heartbeat(run_id, claimed_by))
    try:
        extraction_start = time.perf_counter()
        extracted_text, metadata = await _perform_extraction(run_resume_path)
        parser_name = str(metadata.get("parser", "unknown"))
        _log_event(
            phase="extraction_complete",
            run_id=run_id,
            parser=parser_name,
            duration_ms=int((time.perf_counter() - extraction_start) * 1000),
            chars=len(extracted_text),
        )

        persist_start = time.perf_counter()
        await _insert_resume_document(run_id, run_user_id, run_resume_path, extracted_text, metadata)
        _log_event(
            phase="persist_complete",
            run_id=run_id,
            parser=parser_name,
            duration_ms=int((time.perf_counter() - persist_start) * 1000),
        )

        finalize_start = time.perf_counter()
        await _set_run_ready_for_generation(run_id, claimed_by)
        _log_event(
            phase="finalize_complete",
            run_id=run_id,
            parser=parser_name,
            duration_ms=int((time.perf_counter() - finalize_start) * 1000),
        )
        _log_event(
            phase="extract_success",
            run_id=run_id,
            parser=parser_name,
            duration_ms=int((time.perf_counter() - request_start) * 1000),
            attempt_count=attempt_count,
            queue_wait_ms=queue_wait_ms,
        )
        _log_extraction_summary(
            run_id=run_id,
            status="success",
            parser=parser_name,
            duration_ms=int((time.perf_counter() - request_start) * 1000),
            attempt_count=attempt_count,
            queue_wait_ms=queue_wait_ms,
        )
    except Exception as error:  # pragma: no cover
        _capture_exception(error, phase="process_claimed_run", run_id=run_id, parser=parser_name)
        error_code = error.code if isinstance(error, HttpError) else ERR_PARSE_ERROR
        error_message = str(error) if str(error) else "Unknown extraction failure"
        failure_bucket = _failure_bucket(error_code)
        _log_event(
            phase="extract_error",
            run_id=run_id,
            parser=parser_name,
            duration_ms=int((time.perf_counter() - request_start) * 1000),
            error_code=error_code,
            error_message=error_message,
            failure_bucket=failure_bucket,
            attempt_count=attempt_count,
            queue_wait_ms=queue_wait_ms,
        )
        try:
            if _is_retryable_error(error, error_code) and attempt_count < WORKER_MAX_ATTEMPTS:
                retry_start = time.perf_counter()
                await _requeue_run(
                    run_id,
                    claimed_by,
                    error_code,
                    error_message,
                    WORKER_RETRY_DELAY_SECONDS,
                )
                _log_event(
                    phase="retry_requeued",
                    run_id=run_id,
                    parser=parser_name,
                    duration_ms=int((time.perf_counter() - retry_start) * 1000),
                    error_code=error_code,
                    failure_bucket=failure_bucket,
                    attempt_count=attempt_count,
                    next_attempt=attempt_count + 1,
                    queue_wait_ms=queue_wait_ms,
                )
                _log_extraction_summary(
                    run_id=run_id,
                    status="requeued",
                    parser=parser_name,
                    duration_ms=int((time.perf_counter() - request_start) * 1000),
                    attempt_count=attempt_count,
                    queue_wait_ms=queue_wait_ms,
                    error_code=error_code,
                    failure_bucket=failure_bucket,
                )
            else:
                failed_start = time.perf_counter()
                await _set_run_failed(run_id, claimed_by, error_code, error_message)
                _log_event(
                    phase="failed_status_set",
                    run_id=run_id,
                    parser=parser_name,
                    duration_ms=int((time.perf_counter() - failed_start) * 1000),
                    error_code=error_code,
                    failure_bucket=failure_bucket,
                    attempt_count=attempt_count,
                    queue_wait_ms=queue_wait_ms,
                )
                _log_extraction_summary(
                    run_id=run_id,
                    status="failed",
                    parser=parser_name,
                    duration_ms=int((time.perf_counter() - request_start) * 1000),
                    attempt_count=attempt_count,
                    queue_wait_ms=queue_wait_ms,
                    error_code=error_code,
                    failure_bucket=failure_bucket,
                )
        except Exception:
            _log_event(
                phase="failed_status_set_error",
                run_id=run_id,
                parser=parser_name,
                error_code=error_code,
                failure_bucket=failure_bucket,
                attempt_count=attempt_count,
                queue_wait_ms=queue_wait_ms,
            )
        if isinstance(error, HttpError):
            raise error
        raise HttpError(status=500, code=error_code, message=error_message)
    finally:
        heartbeat_task.cancel()
        with suppress(asyncio.CancelledError):
            await heartbeat_task


async def _run_worker_once() -> bool:
    claimed = await _claim_next_run()
    if claimed is None:
        return False

    run_id_raw = claimed.get("id")
    run_user_id_raw = claimed.get("user_id")
    run_resume_path = claimed.get("resume_path")
    run_attempt_count_raw = claimed.get("extraction_attempt_count")
    run_created_at_raw = claimed.get("created_at")

    if (
        not isinstance(run_id_raw, str)
        or not isinstance(run_user_id_raw, str)
        or not isinstance(run_resume_path, str)
        or not isinstance(run_attempt_count_raw, int)
    ):
        raise HttpError(
            status=500,
            code=ERR_INVALID_RUN_ROW,
            message="Claim-next row is missing id, user_id, resume_path, or extraction_attempt_count",
        )

    run_id = UUID(run_id_raw)
    queue_wait_ms = _compute_queue_wait_ms(run_created_at_raw if isinstance(run_created_at_raw, str) else None)
    _log_event(
        phase="worker_claimed",
        run_id=run_id,
        attempt_count=run_attempt_count_raw,
        queue_wait_ms=queue_wait_ms,
    )
    try:
        await _process_claimed_run(
            run_id,
            run_user_id_raw,
            run_resume_path,
            WORKER_CLAIM_OWNER,
            run_attempt_count_raw,
            queue_wait_ms,
        )
    except HttpError:
        # Failure details and run status are handled in _process_claimed_run.
        pass
    return True


async def _run_generation_worker_once() -> bool:
    claim_start = time.perf_counter()
    row = await _claim_next_generate_run()
    if row is None:
        return False

    run_id_raw = row.get("id")
    request_id = row.get("request_id")
    attempt_count = row.get("generation_attempt_count")
    run_updated_at = row.get("updated_at")
    if not isinstance(run_id_raw, str) or not isinstance(request_id, str) or not isinstance(attempt_count, int):
        raise HttpError(
            status=500,
            code=ERR_INVALID_RUN_ROW,
            message="Claimed generate run is missing id, request_id, or generation_attempt_count",
        )
    run_id = UUID(run_id_raw)
    queue_wait_ms = _compute_queue_wait_ms(run_updated_at if isinstance(run_updated_at, str) else None)

    _log_event(
        phase="generate_claim_complete",
        run_id=run_id,
        duration_ms=int((time.perf_counter() - claim_start) * 1000),
        attempt_count=attempt_count,
        queue_wait_ms=queue_wait_ms,
    )

    try:
        await _process_claimed_generate_run(
            run_id,
            request_id,
            GENERATION_WORKER_CLAIM_OWNER,
            attempt_count,
            queue_wait_ms,
        )
    except HttpError:
        pass
    return True


async def _run_pdf_worker_once() -> bool:
    claim_start = time.perf_counter()
    row = await _claim_next_pdf_run()
    if row is None:
        return False

    run_id_raw = row.get("id")
    attempt_count = row.get("pdf_attempt_count")
    run_updated_at = row.get("updated_at")
    if not isinstance(run_id_raw, str) or not isinstance(attempt_count, int):
        raise HttpError(
            status=500,
            code=ERR_INVALID_RUN_ROW,
            message="Claimed pdf run is missing id or pdf_attempt_count",
        )
    run_id = UUID(run_id_raw)
    queue_wait_ms = _compute_queue_wait_ms(run_updated_at if isinstance(run_updated_at, str) else None)

    _log_event(
        phase="pdf_claim_complete",
        run_id=run_id,
        duration_ms=int((time.perf_counter() - claim_start) * 1000),
        attempt_count=attempt_count,
        queue_wait_ms=queue_wait_ms,
    )

    try:
        await _process_claimed_pdf_run(run_id, attempt_count, queue_wait_ms)
    except HttpError:
        pass
    return True


async def run_worker_forever() -> None:
    wake_event = asyncio.Event()
    realtime_task: asyncio.Task[None] | None = None
    worker_presence_task: asyncio.Task[None] | None = None
    idle_cycles = 0
    idle_log_every_cycles = max(1, int(60 / max(WORKER_POLL_INTERVAL_SECONDS, 0.1)))
    last_queue_snapshot_at = 0.0
    last_stale_reset_at = 0.0
    if WORKER_REALTIME_ENABLED:
        realtime_task = asyncio.create_task(_run_realtime_wakeup_loop(wake_event))
    await _register_worker_presence()
    worker_presence_task = asyncio.create_task(_run_worker_presence_heartbeat())

    _log_event(
        phase="worker_start",
        worker_id=WORKER_CLAIM_OWNER,
        poll_interval_seconds=WORKER_POLL_INTERVAL_SECONDS,
        realtime_enabled=WORKER_REALTIME_ENABLED,
        lease_seconds=WORKER_LEASE_SECONDS,
        run_heartbeat_interval_seconds=RUN_HEARTBEAT_INTERVAL_SECONDS,
        worker_heartbeat_interval_seconds=WORKER_HEARTBEAT_INTERVAL_SECONDS,
        retry_delay_seconds=WORKER_RETRY_DELAY_SECONDS,
        max_attempts=WORKER_MAX_ATTEMPTS,
        queue_snapshot_interval_seconds=WORKER_QUEUE_SNAPSHOT_INTERVAL_SECONDS,
        active_worker_window_seconds=WORKER_ACTIVE_WINDOW_SECONDS,
    )
    try:
        while True:
            try:
                processed_any = False
                while True:
                    processed = await _run_worker_once()
                    generation_processed = False
                    pdf_processed = False
                    if GENERATION_WORKER_ENABLED:
                        generation_processed = await _run_generation_worker_once()
                    if PDF_WORKER_ENABLED:
                        pdf_processed = await _run_pdf_worker_once()
                    if not processed and not generation_processed and not pdf_processed:
                        break
                    processed_any = True

                if processed_any:
                    idle_cycles = 0
                    continue

                idle_cycles += 1
                if idle_cycles % idle_log_every_cycles == 0:
                    _log_event(
                        phase="worker_idle",
                        idle_cycles=idle_cycles,
                        idle_seconds=int(idle_cycles * WORKER_POLL_INTERVAL_SECONDS),
                    )

                now = time.monotonic()
                if now - last_queue_snapshot_at >= WORKER_QUEUE_SNAPSHOT_INTERVAL_SECONDS:
                    snapshot = await _get_queue_snapshot()
                    _log_event(phase="worker_queue_snapshot", **snapshot)
                    last_queue_snapshot_at = now

                last_stale_reset_at = await _run_stale_reset_if_due(last_stale_reset_at)

                try:
                    await asyncio.wait_for(wake_event.wait(), timeout=WORKER_POLL_INTERVAL_SECONDS)
                    wake_event.clear()
                except asyncio.TimeoutError:
                    pass
            except HttpError as error:
                _log_event(phase="worker_error", error_code=error.code, error_message=error.message)
                _capture_exception(error, phase="worker_loop")
                await asyncio.sleep(WORKER_POLL_INTERVAL_SECONDS)
            except Exception as error:  # pragma: no cover
                _log_event(phase="worker_error", error_code=ERR_PARSE_ERROR, error_message=str(error))
                _capture_exception(error, phase="worker_loop")
                await asyncio.sleep(WORKER_POLL_INTERVAL_SECONDS)
    finally:
        if worker_presence_task:
            worker_presence_task.cancel()
            with suppress(asyncio.CancelledError):
                await worker_presence_task
        with suppress(Exception):
            await _release_worker_presence()
        _log_event(phase="worker_stop", worker_id=WORKER_CLAIM_OWNER)
        if realtime_task:
            realtime_task.cancel()
            with suppress(asyncio.CancelledError):
                await realtime_task


def _build_realtime_websocket_url() -> str:
    _assert_config()
    assert SUPABASE_URL is not None
    assert SUPABASE_SERVICE_ROLE_KEY is not None

    parsed = urlparse(SUPABASE_URL)
    scheme = "wss" if parsed.scheme == "https" else "ws"
    query = urlencode({"apikey": SUPABASE_SERVICE_ROLE_KEY, "vsn": "1.0.0"})
    return urlunparse((scheme, parsed.netloc, "/realtime/v1/websocket", "", query, ""))


def _is_resume_runs_queueable_change(message: dict[str, Any]) -> bool:
    if message.get("event") != "postgres_changes":
        return False
    payload = message.get("payload")
    if not isinstance(payload, dict):
        return False
    table = payload.get("table")
    if table != "resume_runs":
        return False
    data = payload.get("data")
    if not isinstance(data, dict):
        return False
    record = data.get("record")
    if not isinstance(record, dict):
        return False
    return record.get("status") in {"queued", "queued_generate", "queued_pdf"}


async def _run_realtime_wakeup_loop(wake_event: asyncio.Event) -> None:
    while True:
        try:
            await _run_realtime_session(wake_event)
        except Exception as error:  # pragma: no cover
            _log_event(
                phase="worker_realtime_error",
                error_code=ERR_DOWNLOAD_FAILED,
                error_message=str(error),
            )
            _capture_exception(error, phase="worker_realtime")
            wake_event.set()
            await asyncio.sleep(WORKER_REALTIME_RECONNECT_SECONDS)


async def _run_realtime_session(wake_event: asyncio.Event) -> None:
    channel_topic = f"realtime:{SUPABASE_DB_SCHEMA}:resume_runs"
    join_payload = {
        "topic": channel_topic,
        "event": "phx_join",
        "payload": {
            "config": {
                "broadcast": {"ack": False, "self": False},
                "presence": {"key": ""},
                "postgres_changes": [
                    {
                        "event": "INSERT",
                        "schema": SUPABASE_DB_SCHEMA,
                        "table": "resume_runs",
                        "filter": "status=eq.queued",
                    },
                    {
                        "event": "UPDATE",
                        "schema": SUPABASE_DB_SCHEMA,
                        "table": "resume_runs",
                        "filter": "status=eq.queued",
                    },
                    {
                        "event": "INSERT",
                        "schema": SUPABASE_DB_SCHEMA,
                        "table": "resume_runs",
                        "filter": "status=eq.queued_generate",
                    },
                    {
                        "event": "UPDATE",
                        "schema": SUPABASE_DB_SCHEMA,
                        "table": "resume_runs",
                        "filter": "status=eq.queued_generate",
                    },
                    {
                        "event": "INSERT",
                        "schema": SUPABASE_DB_SCHEMA,
                        "table": "resume_runs",
                        "filter": "status=eq.queued_pdf",
                    },
                    {
                        "event": "UPDATE",
                        "schema": SUPABASE_DB_SCHEMA,
                        "table": "resume_runs",
                        "filter": "status=eq.queued_pdf",
                    },
                ],
            },
            "access_token": SUPABASE_SERVICE_ROLE_KEY,
        },
        "ref": "1",
    }
    heartbeat_ref = 1
    _log_event(phase="worker_realtime_connecting")

    async with websockets.connect(_build_realtime_websocket_url()) as websocket:
        await websocket.send(json.dumps(join_payload))
        _log_event(phase="worker_realtime_connected")

        async def heartbeat_loop() -> None:
            nonlocal heartbeat_ref
            while True:
                await asyncio.sleep(WORKER_REALTIME_HEARTBEAT_SECONDS)
                heartbeat_ref += 1
                heartbeat = {
                    "topic": "phoenix",
                    "event": "heartbeat",
                    "payload": {},
                    "ref": str(heartbeat_ref),
                }
                await websocket.send(json.dumps(heartbeat))

        heartbeat_task = asyncio.create_task(heartbeat_loop())
        try:
            while True:
                raw = await websocket.recv()
                try:
                    message = json.loads(raw)
                except json.JSONDecodeError:
                    continue
                if _is_resume_runs_queueable_change(message):
                    _log_event(phase="worker_realtime_wakeup")
                    wake_event.set()
        finally:
            heartbeat_task.cancel()


@app.exception_handler(HttpError)
async def http_error_handler(_, exc: HttpError) -> JSONResponse:
    return JSONResponse(
        status_code=exc.status,
        content={
            "ok": False,
            "error_code": exc.code,
            "error_message": exc.message,
        },
    )


@app.get("/health")
async def health() -> dict[str, bool]:
    return {"ok": True}


@app.post("/extract")
async def extract(
    payload: ExtractRequest,
    x_extract_token: str | None = Header(default=None, alias="X-Extract-Token"),
) -> dict[str, Any]:
    _assert_extract_auth(x_extract_token)
    run_id = payload.run_id
    _log_event(phase="extract_start", run_id=run_id)

    claim_start = time.perf_counter()
    run_row = await _claim_run(run_id)
    _log_event(
        phase="claim_complete",
        run_id=run_id,
        duration_ms=int((time.perf_counter() - claim_start) * 1000),
    )

    run_user_id = run_row.get("user_id")
    run_resume_path = run_row.get("resume_path")
    run_attempt_count = run_row.get("extraction_attempt_count")
    run_created_at = run_row.get("created_at")
    if (
        not isinstance(run_user_id, str)
        or not isinstance(run_resume_path, str)
        or not isinstance(run_attempt_count, int)
    ):
        raise HttpError(
            status=500,
            code=ERR_INVALID_RUN_ROW,
            message="Claimed run is missing user_id, resume_path, or extraction_attempt_count",
        )

    queue_wait_ms = _compute_queue_wait_ms(run_created_at if isinstance(run_created_at, str) else None)
    _log_event(
        phase="api_claimed",
        run_id=run_id,
        attempt_count=run_attempt_count,
        queue_wait_ms=queue_wait_ms,
    )

    await _process_claimed_run(
        run_id,
        run_user_id,
        run_resume_path,
        API_CLAIM_OWNER,
        run_attempt_count,
        queue_wait_ms,
    )
    return {"ok": True, "run_id": str(run_id), "status": "queued_generate"}
