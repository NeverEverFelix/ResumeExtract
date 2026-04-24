import asyncio
import json
import unittest
from io import BytesIO
from unittest.mock import AsyncMock, patch
from uuid import uuid4

from docx import Document
from fastapi.testclient import TestClient

import src.main as main


class MainTests(unittest.TestCase):
    def setUp(self) -> None:
        self.client = TestClient(main.app)
        self.original_extract_auth_token = main.EXTRACT_AUTH_TOKEN
        self.original_api_claim_owner = main.API_CLAIM_OWNER
        self.original_worker_claim_owner = main.WORKER_CLAIM_OWNER
        main.EXTRACT_AUTH_TOKEN = "test-extract-token"
        main.API_CLAIM_OWNER = "api:test-owner"
        main.WORKER_CLAIM_OWNER = "worker:test-owner"

    def tearDown(self) -> None:
        main.EXTRACT_AUTH_TOKEN = self.original_extract_auth_token
        main.API_CLAIM_OWNER = self.original_api_claim_owner
        main.WORKER_CLAIM_OWNER = self.original_worker_claim_owner

    def test_health_returns_ok(self) -> None:
        response = self.client.get("/health")
        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.json(), {"ok": True})

    def test_extract_success_contract_flow(self) -> None:
        run_id = uuid4()
        insert_mock = AsyncMock()
        set_ready_mock = AsyncMock()

        with (
            patch.object(
                main,
                "_claim_run",
                AsyncMock(
                    return_value={
                        "id": str(run_id),
                        "user_id": "user-123",
                        "resume_path": "user-123/resume.pdf",
                        "extraction_attempt_count": 1,
                        "created_at": "2026-04-22T18:00:00+00:00",
                    }
                ),
            ),
            patch.object(main, "_perform_extraction", AsyncMock(return_value=("Extracted text", {"parser": "fake"}))),
            patch.object(main, "_insert_resume_document", insert_mock),
            patch.object(main, "_set_run_ready_for_generation", set_ready_mock),
        ):
            response = self.client.post(
                "/extract",
                json={"run_id": str(run_id)},
                headers={"X-Extract-Token": "test-extract-token"},
            )

        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.json(), {"ok": True, "run_id": str(run_id), "status": "queued_generate"})
        insert_mock.assert_awaited_once_with(
            run_id, "user-123", "user-123/resume.pdf", "Extracted text", {"parser": "fake"}
        )
        set_ready_mock.assert_awaited_once_with(run_id, "api:test-owner")

    def test_insert_resume_document_uses_upsert_for_retry_idempotency(self) -> None:
        with patch.object(main, "_supabase_fetch", AsyncMock()) as fetch_mock:
            asyncio.run(
                main._insert_resume_document(
                    run_id=uuid4(),
                    run_user_id="user-123",
                    run_resume_path="user-123/resume.pdf",
                    extracted_text="text",
                    metadata={"parser": "fake"},
                )
            )

        fetch_mock.assert_awaited_once()
        kwargs = fetch_mock.await_args.kwargs
        self.assertEqual(kwargs["path"], "resume_documents?on_conflict=run_id")
        self.assertEqual(kwargs["method"], "POST")
        self.assertEqual(kwargs["prefer"], "resolution=merge-duplicates,return=minimal")

    def test_extract_returns_409_when_not_queued(self) -> None:
        with patch.object(
            main,
            "_claim_run",
            AsyncMock(side_effect=main.HttpError(409, "run_not_queued", "Run is not in queued state or does not exist")),
        ):
            response = self.client.post(
                "/extract",
                json={"run_id": str(uuid4())},
                headers={"X-Extract-Token": "test-extract-token"},
            )

        self.assertEqual(response.status_code, 409)
        self.assertEqual(response.json().get("error_code"), "run_not_queued")

    def test_claim_next_run_uses_rpc_contract(self) -> None:
        row = {
            "id": str(uuid4()),
            "status": "extracting",
            "user_id": "user-123",
            "resume_path": "user-123/resume.pdf",
            "extraction_attempt_count": 1,
            "created_at": "2026-04-22T18:00:00+00:00",
        }
        with patch.object(main, "_supabase_fetch", AsyncMock(return_value=[row])) as fetch_mock:
            claimed = asyncio.run(main._claim_next_run())

        self.assertEqual(claimed, row)
        fetch_mock.assert_awaited_once_with(
            path="rpc/claim_next_resume_run",
            method="POST",
            body={"p_claimed_by": "worker:test-owner", "p_lease_seconds": main.WORKER_LEASE_SECONDS},
        )

    def test_claim_run_stamps_claim_timestamps(self) -> None:
        run_id = uuid4()
        row = {
            "id": str(run_id),
            "status": "extracting",
            "user_id": "user-123",
            "resume_path": "user-123/resume.pdf",
            "extraction_attempt_count": 1,
            "created_at": "2026-04-22T18:00:00+00:00",
        }
        with patch.object(main, "_supabase_fetch", AsyncMock(return_value=[row])) as fetch_mock:
            claimed = asyncio.run(main._claim_run(run_id))

        self.assertEqual(claimed, row)
        kwargs = fetch_mock.await_args.kwargs
        self.assertEqual(kwargs["method"], "POST")
        self.assertEqual(kwargs["path"], "rpc/claim_resume_run")
        self.assertEqual(kwargs["body"], {"p_run_id": str(run_id), "p_claimed_by": "api:test-owner"})

    def test_heartbeat_claimed_run_updates_timestamp(self) -> None:
        run_id = uuid4()
        with patch.object(main, "_supabase_fetch", AsyncMock(return_value=[{"id": str(run_id)}])) as fetch_mock:
            asyncio.run(main._heartbeat_claimed_run(run_id, "worker:test-owner"))

        kwargs = fetch_mock.await_args.kwargs
        self.assertEqual(kwargs["method"], "PATCH")
        self.assertIn(f"resume_runs?id=eq.{run_id}&status=eq.extracting", kwargs["path"])
        self.assertIn("extraction_claimed_by=eq.worker:test-owner", kwargs["path"])
        self.assertIsInstance(kwargs["body"]["extraction_heartbeat_at"], str)

    def test_claim_next_run_returns_none_when_no_queued_rows(self) -> None:
        with patch.object(main, "_supabase_fetch", AsyncMock(return_value=[])):
            claimed = asyncio.run(main._claim_next_run())
        self.assertIsNone(claimed)

    def test_claim_next_run_rejects_invalid_rpc_payload(self) -> None:
        with patch.object(main, "_supabase_fetch", AsyncMock(return_value={"id": "not-a-list"})):
            with self.assertRaises(main.HttpError) as exc:
                asyncio.run(main._claim_next_run())
        self.assertEqual(exc.exception.code, main.ERR_INVALID_RUN_ROW)

    def test_run_worker_once_noop_when_no_claim(self) -> None:
        with patch.object(main, "_claim_next_run", AsyncMock(return_value=None)):
            processed = asyncio.run(main._run_worker_once())
        self.assertFalse(processed)

    def test_run_worker_once_processes_claimed_row(self) -> None:
        run_id = uuid4()
        with (
            patch.object(
                main,
                "_claim_next_run",
                AsyncMock(
                    return_value={
                        "id": str(run_id),
                        "user_id": "user-123",
                        "resume_path": "user-123/resume.pdf",
                        "extraction_attempt_count": 2,
                        "created_at": "2026-04-22T18:00:00+00:00",
                    }
                ),
            ),
            patch.object(main, "_process_claimed_run", AsyncMock()) as process_mock,
        ):
            processed = asyncio.run(main._run_worker_once())

        self.assertTrue(processed)
        process_args = process_mock.await_args.args
        self.assertEqual(process_args[:5], (run_id, "user-123", "user-123/resume.pdf", "worker:test-owner", 2))
        self.assertIsInstance(process_args[5], int)

    def test_run_worker_once_rejects_invalid_claimed_row(self) -> None:
        with patch.object(main, "_claim_next_run", AsyncMock(return_value={"id": "abc"})):
            with self.assertRaises(main.HttpError) as exc:
                asyncio.run(main._run_worker_once())
        self.assertEqual(exc.exception.code, main.ERR_INVALID_RUN_ROW)

    def test_extract_marks_failed_on_exception(self) -> None:
        run_id = uuid4()
        set_failed_mock = AsyncMock()

        with (
            patch.object(
                main,
                "_claim_run",
                AsyncMock(
                    return_value={
                        "id": str(run_id),
                        "user_id": "user-123",
                        "resume_path": "user-123/resume.pdf",
                        "extraction_attempt_count": 1,
                        "created_at": "2026-04-22T18:00:00+00:00",
                    }
                ),
            ),
            patch.object(main, "_perform_extraction", AsyncMock(side_effect=ValueError("boom"))),
            patch.object(main, "_set_run_failed", set_failed_mock),
            patch.object(main, "_requeue_run", AsyncMock()),
            patch.object(main, "_set_run_ready_for_generation", AsyncMock()),
        ):
            response = self.client.post(
                "/extract",
                json={"run_id": str(run_id)},
                headers={"X-Extract-Token": "test-extract-token"},
            )

        self.assertEqual(response.status_code, 500)
        payload = response.json()
        self.assertEqual(payload.get("error_code"), main.ERR_PARSE_ERROR)
        set_failed_mock.assert_awaited_once()
        args = set_failed_mock.await_args.args
        self.assertEqual(str(args[0]), str(run_id))
        self.assertEqual(args[1], "api:test-owner")
        self.assertEqual(args[2], main.ERR_PARSE_ERROR)
        self.assertEqual(args[3], "boom")

    def test_process_claimed_run_requeues_retryable_error_before_max_attempts(self) -> None:
        run_id = uuid4()
        with (
            patch.object(main, "_run_claim_heartbeat", AsyncMock()),
            patch.object(
                main,
                "_perform_extraction",
                AsyncMock(side_effect=main.HttpError(502, main.ERR_DOWNLOAD_FAILED, "temporary download issue")),
            ),
            patch.object(main, "_requeue_run", AsyncMock()) as requeue_mock,
            patch.object(main, "_set_run_failed", AsyncMock()) as failed_mock,
        ):
            with self.assertRaises(main.HttpError) as exc:
                asyncio.run(
                    main._process_claimed_run(
                        run_id,
                        "user-123",
                        "user-123/resume.pdf",
                        "worker:test-owner",
                        1,
                        500,
                    )
                )

        self.assertEqual(exc.exception.code, main.ERR_DOWNLOAD_FAILED)
        requeue_mock.assert_awaited_once_with(
            run_id,
            "worker:test-owner",
            main.ERR_DOWNLOAD_FAILED,
            "temporary download issue",
            main.WORKER_RETRY_DELAY_SECONDS,
        )
        failed_mock.assert_not_awaited()

    def test_process_claimed_run_fails_retryable_error_at_max_attempts(self) -> None:
        run_id = uuid4()
        with (
            patch.object(main, "_run_claim_heartbeat", AsyncMock()),
            patch.object(
                main,
                "_perform_extraction",
                AsyncMock(side_effect=main.HttpError(502, main.ERR_DOWNLOAD_FAILED, "temporary download issue")),
            ),
            patch.object(main, "_requeue_run", AsyncMock()) as requeue_mock,
            patch.object(main, "_set_run_failed", AsyncMock()) as failed_mock,
        ):
            with self.assertRaises(main.HttpError):
                asyncio.run(
                    main._process_claimed_run(
                        run_id,
                        "user-123",
                        "user-123/resume.pdf",
                        "worker:test-owner",
                        main.WORKER_MAX_ATTEMPTS,
                        500,
                    )
                )

        requeue_mock.assert_not_awaited()
        failed_mock.assert_awaited_once_with(
            run_id,
            "worker:test-owner",
            main.ERR_DOWNLOAD_FAILED,
            "temporary download issue",
        )

    def test_compute_queue_wait_ms_returns_none_for_invalid_timestamp(self) -> None:
        self.assertIsNone(main._compute_queue_wait_ms("not-a-timestamp"))

    def test_resume_runs_count_path_builds_query(self) -> None:
        self.assertEqual(
            main._resume_runs_count_path("status=eq.queued", "created_at=lte.2026-04-22T18%3A00%3A00%2B00%3A00"),
            "resume_runs?select=id&status=eq.queued&created_at=lte.2026-04-22T18%3A00%3A00%2B00%3A00",
        )

    def test_get_queue_snapshot_aggregates_counts(self) -> None:
        with patch.object(main, "_supabase_count", AsyncMock(side_effect=[10, 4, 2, 1, 3, 1, 2])):
            snapshot = asyncio.run(main._get_queue_snapshot())

        self.assertEqual(
            snapshot,
            {
                "queued_total": 10,
                "queued_over_1m": 4,
                "queued_over_5m": 2,
                "queued_over_15m": 1,
                "extracting_total": 3,
                "stale_extracting_total": 1,
                "active_workers": 2,
            },
        )

    def test_register_worker_presence_uses_upsert(self) -> None:
        with patch.object(main, "_supabase_fetch", AsyncMock()) as fetch_mock:
            asyncio.run(main._register_worker_presence())

        kwargs = fetch_mock.await_args.kwargs
        self.assertEqual(kwargs["path"], "extract_worker_heartbeats?on_conflict=worker_id")
        self.assertEqual(kwargs["method"], "POST")
        self.assertEqual(kwargs["body"]["worker_id"], "worker:test-owner")
        self.assertEqual(kwargs["body"]["role"], "extractor")
        self.assertEqual(kwargs["prefer"], "resolution=merge-duplicates,return=minimal")

    def test_release_worker_presence_deletes_row(self) -> None:
        with patch.object(main, "_supabase_fetch", AsyncMock()) as fetch_mock:
            asyncio.run(main._release_worker_presence())

        kwargs = fetch_mock.await_args.kwargs
        self.assertEqual(kwargs["method"], "DELETE")
        self.assertIn("extract_worker_heartbeats?worker_id=eq.worker:test-owner", kwargs["path"])
        self.assertEqual(kwargs["prefer"], "return=minimal")

    def test_reset_stale_runs_uses_rpc_contract(self) -> None:
        rows = [{"id": str(uuid4())}]
        with patch.object(main, "_supabase_fetch", AsyncMock(return_value=rows)) as fetch_mock:
            result = asyncio.run(main._reset_stale_runs(limit=25))

        self.assertEqual(result, rows)
        fetch_mock.assert_awaited_once_with(
            path="rpc/reset_stale_resume_runs",
            method="POST",
            body={
                "p_stale_seconds": main.WORKER_LEASE_SECONDS,
                "p_limit": 25,
            },
        )

    def test_run_stale_reset_if_due_resets_and_logs_count(self) -> None:
        with (
            patch.object(main, "_reset_stale_runs", AsyncMock(return_value=[{"id": "a"}, {"id": "b"}])) as reset_mock,
            patch.object(main, "_log_event") as log_mock,
            patch.object(main.time, "monotonic", return_value=120.0),
        ):
            next_mark = asyncio.run(main._run_stale_reset_if_due(0.0))

        self.assertEqual(next_mark, 120.0)
        reset_mock.assert_awaited_once()
        log_mock.assert_called_once_with(phase="worker_stale_reset", reset_count=2)

    def test_failure_bucket_classifies_retryable_infra_errors(self) -> None:
        self.assertEqual(main._failure_bucket(main.ERR_DB_ERROR), "db")
        self.assertEqual(main._failure_bucket(main.ERR_DOWNLOAD_FAILED), "download")

    def test_failure_bucket_classifies_parser_and_system_errors(self) -> None:
        self.assertEqual(main._failure_bucket(main.ERR_PARSE_ERROR), "parse")
        self.assertEqual(main._failure_bucket(main.ERR_MISSING_CONFIG), "system")

    def test_log_event_serializes_extra_fields(self) -> None:
        with patch.object(main.logger, "info") as info_mock:
            main._log_event(
                phase="worker_start",
                worker_id="worker:test-owner",
                lease_seconds=120,
                max_attempts=3,
            )

        payload = json.loads(info_mock.call_args.args[0])
        self.assertEqual(payload["phase"], "worker_start")
        self.assertEqual(payload["worker_id"], "worker:test-owner")
        self.assertEqual(payload["lease_seconds"], 120)
        self.assertEqual(payload["max_attempts"], 3)

    def test_log_extraction_summary_serializes_consistent_shape(self) -> None:
        run_id = uuid4()
        with patch.object(main.logger, "info") as info_mock:
            main._log_extraction_summary(
                run_id=run_id,
                status="requeued",
                parser="pypdf",
                duration_ms=2500,
                attempt_count=2,
                queue_wait_ms=900,
                error_code=main.ERR_DOWNLOAD_FAILED,
                failure_bucket="download",
            )

        payload = json.loads(info_mock.call_args.args[0])
        self.assertEqual(payload["phase"], "extract_summary")
        self.assertEqual(payload["run_id"], str(run_id))
        self.assertEqual(payload["status"], "requeued")
        self.assertEqual(payload["parser"], "pypdf")
        self.assertEqual(payload["duration_ms"], 2500)
        self.assertEqual(payload["attempt_count"], 2)
        self.assertEqual(payload["queue_wait_ms"], 900)
        self.assertEqual(payload["error_code"], main.ERR_DOWNLOAD_FAILED)
        self.assertEqual(payload["failure_bucket"], "download")

    def test_normalize_text_removes_empty_lines_and_extra_spaces(self) -> None:
        raw = "  Jane   Doe  \n\n Senior   Engineer \n\n  Python  "
        self.assertEqual(main._normalize_text(raw), "Jane Doe\nSenior Engineer\nPython")

    def test_normalize_text_normalizes_bullets(self) -> None:
        raw = "• Python\n- SQL\n  * Leadership"
        self.assertEqual(main._normalize_text(raw), "- Python\n- SQL\n- Leadership")

    def test_normalize_text_removes_page_markers_and_consecutive_duplicates(self) -> None:
        raw = "Page 1 of 2\nJane Doe\nJane Doe\n2/2\nSenior Engineer"
        self.assertEqual(main._normalize_text(raw), "Jane Doe\nSenior Engineer")

    def test_normalize_text_removes_weird_whitespace(self) -> None:
        raw = "Jane\u00a0Doe\u200b\nSenior\tEngineer"
        self.assertEqual(main._normalize_text(raw), "Jane Doe\nSenior Engineer")

    def test_normalize_text_reduces_repeated_boilerplate_for_pdf_mode(self) -> None:
        raw = "Jane Doe\nACME Inc.\nExperience\nACME Inc.\nSkills\nACME Inc.\nProjects"
        self.assertEqual(
            main._normalize_text(raw, remove_repeated_boilerplate=True),
            "Jane Doe\nACME Inc.\nExperience\nSkills\nProjects",
        )

    def test_extract_docx_parses_text(self) -> None:
        document = Document()
        document.add_paragraph("Jane Doe")
        document.add_paragraph("Senior Engineer")
        buf = BytesIO()
        document.save(buf)

        text, metadata = main._extract_docx(buf.getvalue())
        self.assertIn("Jane Doe", text)
        self.assertIn("Senior Engineer", text)
        self.assertEqual(metadata.get("parser"), "python-docx")

    def test_perform_extraction_rejects_unsupported_extension(self) -> None:
        with patch.object(main, "_download_resume_bytes", AsyncMock(return_value=b"dummy")):
            with self.assertRaises(main.HttpError) as exc:
                asyncio.run(main._perform_extraction("user-123/resume.txt"))
        self.assertEqual(exc.exception.code, main.ERR_UNSUPPORTED_FILE_TYPE)

    def test_extract_pdf_uses_ocr_fallback_when_text_is_empty(self) -> None:
        class FakePage:
            def extract_text(self):
                return ""

        class FakeReader:
            def __init__(self, _):
                self.pages = [FakePage(), FakePage()]

        with (
            patch.object(main, "PdfReader", FakeReader),
            patch.object(main, "_extract_pdf_with_ocr", return_value=("OCR text", {"parser": "pypdf+tesseract"})) as ocr_mock,
        ):
            text, metadata = main._extract_pdf(b"%PDF-1.4 dummy")

        self.assertEqual(text, "OCR text")
        self.assertEqual(metadata["parser"], "pypdf+tesseract")
        ocr_mock.assert_called_once()

    def test_extract_pdf_returns_ocr_unavailable_when_tesseract_missing(self) -> None:
        with patch.object(main.shutil, "which", return_value=None):
            with self.assertRaises(main.HttpError) as exc:
                main._extract_pdf_with_ocr(b"%PDF-1.4 dummy", 1)
        self.assertEqual(exc.exception.code, main.ERR_OCR_UNAVAILABLE)

    def test_extract_rejects_missing_auth_header(self) -> None:
        response = self.client.post("/extract", json={"run_id": str(uuid4())})
        self.assertEqual(response.status_code, 401)
        self.assertEqual(response.json().get("error_code"), main.ERR_UNAUTHORIZED)

    def test_extract_rejects_wrong_auth_header(self) -> None:
        response = self.client.post(
            "/extract",
            json={"run_id": str(uuid4())},
            headers={"X-Extract-Token": "wrong-token"},
        )
        self.assertEqual(response.status_code, 401)
        self.assertEqual(response.json().get("error_code"), main.ERR_UNAUTHORIZED)

    def test_validate_content_type_rejects_mismatch(self) -> None:
        with self.assertRaises(main.HttpError) as exc:
            main._validate_content_type("user-123/resume.pdf", "text/plain")
        self.assertEqual(exc.exception.code, main.ERR_UNSUPPORTED_CONTENT_TYPE)

    def test_validate_file_size_rejects_large_payload(self) -> None:
        with patch.object(main, "MAX_RESUME_FILE_SIZE_BYTES", 10), patch.object(main, "MAX_RESUME_FILE_SIZE_MB", 0):
            with self.assertRaises(main.HttpError) as exc:
                main._validate_file_size("11", 5)
        self.assertEqual(exc.exception.code, main.ERR_FILE_TOO_LARGE)

    def test_download_resume_maps_timeout_error(self) -> None:
        class FakeTimeoutClient:
            async def __aenter__(self):
                return self

            async def __aexit__(self, exc_type, exc, tb):
                return False

            async def get(self, *args, **kwargs):
                raise main.httpx.TimeoutException("timeout")

        with (
            patch.object(main.httpx, "AsyncClient", return_value=FakeTimeoutClient()),
            patch.object(main, "SUPABASE_URL", "https://example.supabase.co"),
            patch.object(main, "SUPABASE_SERVICE_ROLE_KEY", "key"),
            patch.object(main, "SUPABASE_STORAGE_BUCKET", "Resumes"),
        ):
            with self.assertRaises(main.HttpError) as exc:
                asyncio.run(main._download_resume_bytes("user-123/resume.pdf"))
        self.assertEqual(exc.exception.code, main.ERR_DOWNLOAD_FAILED)

    def test_build_realtime_websocket_url_uses_ws_scheme_and_path(self) -> None:
        with (
            patch.object(main, "SUPABASE_URL", "https://example.supabase.co"),
            patch.object(main, "SUPABASE_SERVICE_ROLE_KEY", "service-role-key"),
            patch.object(main, "SUPABASE_STORAGE_BUCKET", "Resumes"),
        ):
            url = main._build_realtime_websocket_url()
        self.assertEqual(
            url,
            "wss://example.supabase.co/realtime/v1/websocket?apikey=service-role-key&vsn=1.0.0",
        )

    def test_is_resume_runs_queued_change_true_for_queued_record(self) -> None:
        message = {
            "event": "postgres_changes",
            "payload": {
                "table": "resume_runs",
                "data": {
                    "record": {"status": "queued"},
                },
            },
        }
        self.assertTrue(main._is_resume_runs_queued_change(message))

    def test_is_resume_runs_queued_change_false_for_non_queued_record(self) -> None:
        message = {
            "event": "postgres_changes",
            "payload": {
                "table": "resume_runs",
                "data": {
                    "record": {"status": "extracting"},
                },
            },
        }
        self.assertFalse(main._is_resume_runs_queued_change(message))


if __name__ == "__main__":
    unittest.main()
